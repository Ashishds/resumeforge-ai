"""
Document Processing Service
Handles file parsing, text extraction, and format conversion
"""

import io
from typing import Tuple
from pypdf import PdfReader
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT


class DocumentExtractor:
    """Extracts text content from various document formats"""
    
    @staticmethod
    def extract_from_pdf(file_bytes: bytes) -> str:
        """
        Extract text from PDF file
        
        Args:
            file_bytes: PDF file content as bytes
            
        Returns:
            str: Extracted text content
        """
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            text_segments = []
            
            for page in reader.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_segments.append(page_text)
            
            return "\n".join(text_segments)
        except Exception as e:
            raise ValueError(f"PDF extraction failed: {str(e)}")
    
    @staticmethod
    def extract_from_docx(file_bytes: bytes) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_bytes: DOCX file content as bytes
            
        Returns:
            str: Extracted text content
        """
        try:
            document_stream = io.BytesIO(file_bytes)
            doc = Document(document_stream)
            
            text_segments = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_segments.append(paragraph.text)
            
            return "\n".join(text_segments)
        except Exception as e:
            raise ValueError(f"DOCX extraction failed: {str(e)}")
    
    @staticmethod
    def extract_from_text(file_bytes: bytes) -> str:
        """
        Extract text from plain text file
        
        Args:
            file_bytes: Text file content as bytes
            
        Returns:
            str: Decoded text content
        """
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception as e:
            return file_bytes.decode("latin-1", errors="ignore")
    
    @staticmethod
    def detect_and_extract(filename: str, file_bytes: bytes) -> Tuple[str, str]:
        """
        Automatically detect file type and extract text
        
        Args:
            filename: Original filename with extension
            file_bytes: File content as bytes
            
        Returns:
            Tuple[str, str]: (file_type, extracted_text)
        """
        filename_lower = filename.lower()
        
        if filename_lower.endswith(".pdf"):
            return "pdf", DocumentExtractor.extract_from_pdf(file_bytes)
        elif filename_lower.endswith(".docx"):
            return "docx", DocumentExtractor.extract_from_docx(file_bytes)
        elif filename_lower.endswith((".txt", ".text")):
            return "txt", DocumentExtractor.extract_from_text(file_bytes)
        else:
            # Attempt text extraction as fallback
            return "unknown", DocumentExtractor.extract_from_text(file_bytes)


class DocumentGenerator:
    """Generates documents in various formats"""
    
    @staticmethod
    def text_to_docx_bytes(text: str) -> bytes:
        """
        Convert plain text to DOCX format
        
        Args:
            text: Text content to convert
            
        Returns:
            bytes: DOCX file as bytes
        """
        doc = Document()
        
        # Add each line as a paragraph
        for line in text.splitlines():
            if line.strip():
                doc.add_paragraph(line)
            else:
                doc.add_paragraph("")  # Empty paragraph for spacing
        
        # Save to bytes
        output_stream = io.BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)
        
        return output_stream.getvalue()
    
    @staticmethod
    def text_to_pdf_bytes(text: str, title: str = "Resume") -> bytes:
        """
        Convert plain text to professionally formatted PDF
        
        Args:
            text: Resume text content
            title: Document title
            
        Returns:
            bytes: PDF file as bytes
        """
        buffer = io.BytesIO()
        
        # Create PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        
        # Custom styles for resume
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor='#1a1a1a',
            spaceAfter=6,
            alignment=TA_LEFT,
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=12,
            textColor='#2c3e50',
            spaceAfter=6,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            textColor='#333333',
            spaceAfter=4,
            alignment=TA_LEFT,
            fontName='Helvetica'
        )
        
        bullet_style = ParagraphStyle(
            'CustomBullet',
            parent=styles['Normal'],
            fontSize=10,
            textColor='#333333',
            spaceAfter=4,
            leftIndent=20,
            fontName='Helvetica'
        )
        
        # Parse and format the resume text
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                elements.append(Spacer(1, 0.1*inch))
                continue
            
            # Skip horizontal dividers
            if line == '---' or line == '___':
                elements.append(Spacer(1, 0.15*inch))
                continue
            
            # Process bold markdown (**text**)
            # Replace **text** with <b>text</b> for ReportLab
            processed_line = line.replace('**', '<b>', 1)
            if '<b>' in processed_line:
                processed_line = processed_line.replace('**', '</b>', 1)
            
            # Detect section headers (all caps or bold headers)
            if line.isupper() and len(line) < 50:
                elements.append(Paragraph(processed_line, heading_style))
            # Detect bold headers (markdown style)
            elif line.startswith('**') and line.endswith('**'):
                clean_line = line.replace('**', '')
                elements.append(Paragraph(clean_line, heading_style))
            # Detect name/title at the top (first non-empty line)
            elif len(elements) == 0 or (len(elements) == 1 and isinstance(elements[0], Spacer)):
                elements.append(Paragraph(processed_line, title_style))
            # Detect bullet points
            elif line.startswith('-') or line.startswith('•'):
                bullet_text = line[1:].strip()
                # Process bold in bullet text
                bullet_text = bullet_text.replace('**', '<b>', 1)
                if '<b>' in bullet_text:
                    bullet_text = bullet_text.replace('**', '</b>', 1)
                elements.append(Paragraph(f"• {bullet_text}", bullet_style))
            else:
                elements.append(Paragraph(processed_line, normal_style))
        
        # Build PDF
        doc.build(elements)
        
        # Get the value of the BytesIO buffer
        buffer.seek(0)
        return buffer.getvalue()


class DocumentValidator:
    """Validates document content and quality"""
    
    @staticmethod
    def validate_resume_content(text: str) -> Tuple[bool, str]:
        """
        Validate that extracted text is suitable for processing
        
        Args:
            text: Extracted text content
            
        Returns:
            Tuple[bool, str]: (is_valid, error_message)
        """
        if not text or not text.strip():
            return False, "Document appears to be empty"
        
        if len(text.strip()) < 100:
            return False, "Document content is too short (minimum 100 characters)"
        
        if len(text.strip()) > 50000:
            return False, "Document content is too long (maximum 50,000 characters)"
        
        # Check for meaningful content (not just random characters)
        words = text.split()
        if len(words) < 20:
            return False, "Document does not contain enough words"
        
        return True, ""
